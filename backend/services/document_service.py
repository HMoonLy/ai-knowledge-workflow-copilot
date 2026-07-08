import json
from io import BytesIO
from pathlib import  Path
from  uuid import uuid4

from pypdf import PdfReader
from docx import Document as DocxDocument
from fastapi import HTTPException, UploadFile


from schemas import DocumentUploadResponse,Document

UPLOAD_DIR=Path('uploads')
DOCUMENTS_FILE=Path('documents.json')


def load_documents_from_file()->dict[int,list[Document]]:
    if not DOCUMENTS_FILE.exists():
        return {}

    raw_text=DOCUMENTS_FILE.read_text(encoding='utf-8')
    raw_data=json.loads(raw_text)
    loaded_documents:dict[int,list[Document]]={}

    for knowledge_base_id_text, documents in raw_data.items():
        knowledge_base_id=int(knowledge_base_id_text)
        loaded_documents[knowledge_base_id]=[]

        for document_data in documents:
            document=Document(
                id=document_data['id'],
                filename=document_data['filename'],
                content_type=document_data['content_type'],
                size=document_data['size'],
                content=document_data['content'],
            )
            loaded_documents[knowledge_base_id].append(document)

    return loaded_documents


def save_documents_to_file()->None:
    data={}

    for knowledge_base_id, documents in document_by_kb.items():
        data[str(knowledge_base_id)]=[]

        for document in documents:
            data[str(knowledge_base_id)].append({
                'id':document.id,
                'filename':document.filename,
                'content_type':document.content_type,
                'size':document.size,
                'content':document.content,
            })

    DOCUMENTS_FILE.write_text(
        json.dumps(data, ensure_ascii=False, indent=2),
        encoding='utf-8',
    )


document_by_kb:dict[int,list[Document]] = load_documents_from_file()


async def save_uploaded_document(file: UploadFile, knowledge_base_id: int, )->DocumentUploadResponse:
    UPLOAD_DIR.mkdir(exist_ok=True)

    file_id=uuid4().hex
    safe_filename=Path(file.filename or 'untitled').name
    saved_path=UPLOAD_DIR/f"{file_id}_{safe_filename}"

    content=await file.read()
    text_content=extract_text_from_file(safe_filename, content)
    if not text_content.strip():
        raise HTTPException(
            status_code=400,
            detail="暂不支持该文件类型，或文件中没有可提取的文字内容",
        )
    saved_path.write_bytes(content)

    document=Document(
        id=len(document_by_kb.get(knowledge_base_id,[]))+1,
        filename=safe_filename,
        content_type=file.content_type or 'application/octet-stream',
        size=len(content),
        content=text_content,
    )
    document_by_kb.setdefault(knowledge_base_id,[]).append(document)
    save_documents_to_file()

    return DocumentUploadResponse(
        id=document.id,
        filename=safe_filename,
        content_type=file.content_type or 'application/octet-stream',
        size=len(content),
        status="uploaded"
    )

def list_documents(knowledge_base_id:int)->list[Document]:
    return document_by_kb.get(knowledge_base_id,[])


def  delete_document(knowledge_base_id:int,document_id:int)->bool:
    documents=document_by_kb.get(knowledge_base_id,[])

    for document in documents:
        if document.id==document_id:
            documents.remove(document)
            save_documents_to_file()
            return True
    return False


def extract_text_from_file(filename: str, content: bytes) -> str:
    filename_lower = filename.lower()

    if filename_lower.endswith(".txt") or filename_lower.endswith(".md"):
        return content.decode("utf-8", errors="ignore")
    if filename_lower.endswith(".pdf"):
        pdf_file=BytesIO(content)
        reader=PdfReader(pdf_file)

        text=""

        for page in reader.pages:
            page_text=page.extractText()

            if page_text:
                text+=page_text+'\n'
        return text

    if filename_lower.endswith(".docx"):
        docx_file = BytesIO(content)
        docx_document = DocxDocument(docx_file)

        text = ""

        for paragraph in docx_document.paragraphs:
            if paragraph.text:
                text += paragraph.text + "\n"

        return text
    return ""

